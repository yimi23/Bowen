#!/usr/bin/env python3
"""
BOWEN Framework Document Engine - Simple and Effective
Just creates documents in any format without bloat
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional

class AdvancedDocumentEngine:
    """
    Simple document creation that actually works
    - Create documents in any format (PDF, DOCX, HTML, TXT)
    - No bullshit themes or overcomplicated features
    - Just create the fucking document the user wants
    """
    
    def __init__(self, output_dir: str = None):
        if output_dir is None:
            output_dir = os.getenv('BOWEN_OUTPUT_DIR', str(Path.home() / "Desktop"))
        
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
    
    def create_document(self, content: str, filename: str, doc_type: str = "docx") -> str:
        """Create any document in any format - simple and effective"""
        try:
            filepath = self.output_dir / f"{filename}.{doc_type}"
            
            if doc_type == "docx":
                return self._create_word_doc(content, filepath)
            elif doc_type == "pdf":
                return self._create_pdf_doc(content, filepath)
            elif doc_type == "html":
                return self._create_html_doc(content, filepath)
            else:  # txt, md, etc.
                return self._create_text_doc(content, filepath)
                
        except Exception as e:
            return f"Document creation failed: {e}"
    
    def _create_word_doc(self, content: str, filepath: Path) -> str:
        """Create Word document"""
        try:
            from docx import Document
            
            doc = Document()
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.strip().startswith('#'):
                        # Heading
                        heading_text = para.strip().lstrip('#').strip()
                        level = min(len(para) - len(para.lstrip('#')), 3)
                        doc.add_heading(heading_text, level)
                    else:
                        # Regular paragraph
                        doc.add_paragraph(para.strip())
            
            doc.save(str(filepath))
            return str(filepath)
        except ImportError:
            # Auto-install if missing
            import subprocess
            subprocess.run(['pip', 'install', 'python-docx'], capture_output=True)
            return self._create_word_doc(content, filepath)
    
    def _create_pdf_doc(self, content: str, filepath: Path) -> str:
        """Create PDF document"""
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import letter
            
            doc = SimpleDocTemplate(str(filepath), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            return str(filepath)
        except ImportError:
            # Auto-install if missing
            import subprocess
            subprocess.run(['pip', 'install', 'reportlab'], capture_output=True)
            return self._create_pdf_doc(content, filepath)
    
    def _create_html_doc(self, content: str, filepath: Path) -> str:
        """Create HTML document"""
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <title>{filepath.stem}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #333; }}
    </style>
</head>
<body>
    <div>{'<br><br>'.join(content.split('\n\n'))}</div>
</body>
</html>"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return str(filepath)
    
    def _create_text_doc(self, content: str, filepath: Path) -> str:
        """Create text document"""
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return str(filepath)